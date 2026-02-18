#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert external raw docs (docx/pdf/html) to structured Markdown."""

import datetime
import hashlib
import os
import re
from pathlib import Path


BASE_DIR = Path("/root/crmeb-java/hxy")
SOURCE_ROOT = BASE_DIR / "external" / "raw"
OUTPUT_ROOT = BASE_DIR / "external" / "structured"


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        while True:
            chunk = f.read(1024 * 1024)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def write_meta_header(out, title: str, source_rel: str, source_type: str, sha256: str) -> None:
    ts = datetime.datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %z")
    out.write(f"# {title}\n\n")
    out.write("## 元数据\n")
    out.write(f"- 来源文件: external/raw/{source_rel}\n")
    out.write(f"- 来源类型: {source_type}\n")
    out.write(f"- SHA256: {sha256}\n")
    out.write(f"- 转换时间: {ts}\n\n")
    out.write("## 提取正文\n")


def convert_docx_to_md(src: Path, dst: Path, source_rel: str) -> bool:
    try:
        from docx import Document

        doc = Document(str(src))
        with dst.open("w", encoding="utf-8") as out:
            write_meta_header(out, src.stem, source_rel, "docx", file_sha256(src))

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = getattr(para.style, "name", "") or ""
                if style_name.startswith("Heading") and style_name[-1:].isdigit():
                    level = int(style_name[-1])
                    out.write(f"{'#' * max(1, min(level, 6))} {text}\n\n")
                else:
                    out.write(f"{text}\n\n")

            for table in doc.tables:
                if not table.rows:
                    continue
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                out.write("| " + " | ".join(header_cells) + " |\n")
                out.write("| " + " | ".join(["---"] * len(header_cells)) + " |\n")
                for row in table.rows[1:]:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    out.write("| " + " | ".join(row_cells) + " |\n")
                out.write("\n")

        print(f"✅ {src} -> {dst}")
        return True
    except ImportError:
        print("❌ 缺少 python-docx，安装: pip install python-docx")
        return False
    except Exception as err:
        print(f"❌ DOCX 转换失败 {src}: {err}")
        return False


def convert_pdf_to_md(src: Path, dst: Path, source_rel: str) -> bool:
    try:
        import pdfplumber

        with pdfplumber.open(str(src)) as pdf, dst.open("w", encoding="utf-8") as out:
            write_meta_header(out, src.stem, source_rel, "pdf", file_sha256(src))
            out.write("以下内容来自 PDF 自动提取，已保留版式换行。\n\n")
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    out.write(text)
                    out.write("\n\n")

        print(f"✅ {src} -> {dst}")
        return True
    except ImportError:
        print("❌ 缺少 pdfplumber，安装: pip install pdfplumber")
        return False
    except Exception as err:
        print(f"❌ PDF 转换失败 {src}: {err}")
        return False


def html_to_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return soup.get_text("\n")
    except Exception:
        text = re.sub(r"<script[\\s\\S]*?</script>", "", html, flags=re.IGNORECASE)
        text = re.sub(r"<style[\\s\\S]*?</style>", "", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", "\n", text)
        return text


def convert_html_to_md(src: Path, dst: Path, source_rel: str) -> bool:
    try:
        raw = src.read_text(encoding="utf-8", errors="ignore")
        text = html_to_text(raw)
        lines = [line.strip() for line in text.splitlines()]
        lines = [line for line in lines if line]

        with dst.open("w", encoding="utf-8") as out:
            write_meta_header(out, src.stem, source_rel, "html", file_sha256(src))
            for line in lines:
                out.write(line)
                out.write("\n\n")

        print(f"✅ {src} -> {dst}")
        return True
    except Exception as err:
        print(f"❌ HTML 转换失败 {src}: {err}")
        return False


def convert_one(src: Path, dst: Path, source_rel: str) -> bool:
    suffix = src.suffix.lower()
    if suffix == ".docx":
        return convert_docx_to_md(src, dst, source_rel)
    if suffix == ".pdf":
        return convert_pdf_to_md(src, dst, source_rel)
    if suffix == ".html":
        return convert_html_to_md(src, dst, source_rel)
    return False


def main() -> None:
    if not SOURCE_ROOT.exists():
        print(f"❌ 来源目录不存在: {SOURCE_ROOT}")
        return

    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

    candidates = []
    for src in SOURCE_ROOT.rglob("*"):
        if src.suffix.lower() in {".docx", ".pdf", ".html"}:
            candidates.append(src)

    print(f"待转换文档数: {len(candidates)}")

    ok = 0
    fail = 0
    for src in sorted(candidates):
        rel = src.relative_to(SOURCE_ROOT)
        dst = OUTPUT_ROOT / rel.with_suffix(".md")
        dst.parent.mkdir(parents=True, exist_ok=True)
        if convert_one(src, dst, str(rel)):
            ok += 1
        else:
            fail += 1

    print("=" * 50)
    print(f"转换完成: 成功 {ok}，失败 {fail}")
    print(f"输出目录: {OUTPUT_ROOT}")
    print("=" * 50)


if __name__ == "__main__":
    main()
